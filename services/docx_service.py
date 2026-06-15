from docx import Document


def generate_docx(
    filename: str,
    query: str,
    report: str
):

    doc = Document()

    doc.add_heading(
        "Research Report",
        level=1
    )

    doc.add_heading(
        query,
        level=2
    )

    doc.add_paragraph(
        report
    )

    doc.save(
        filename
    )

    return filename